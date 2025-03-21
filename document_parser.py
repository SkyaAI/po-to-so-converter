#!/usr/bin/env python3
"""
Document Parser Module for Purchase Order to Sales Order Conversion System

This module provides parsers for different document formats:
- PDF
- Images (using OCR)
- Excel
- Word

It also provides a unified interface for document parsing.
"""

import os
import io
import re
import pandas as pd
import pdfplumber
import pytesseract
from PIL import Image
from pdf2image import convert_from_path, convert_from_bytes
import docx
import cv2
import numpy as np

class DocumentParser:
    """Base class for document parsers"""
    
    def __init__(self):
        self.content = None
        self.file_path = None
        self.file_type = None
    
    def parse(self, file_path):
        """Parse document and return content"""
        self.file_path = file_path
        self._detect_file_type()
        
        if self.file_type == 'pdf':
            return self._parse_pdf()
        elif self.file_type in ['jpg', 'jpeg', 'png', 'tiff', 'bmp']:
            return self._parse_image()
        elif self.file_type in ['xlsx', 'xls']:
            return self._parse_excel()
        elif self.file_type in ['docx', 'doc']:
            return self._parse_word()
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")
    
    def _detect_file_type(self):
        """Detect file type from extension"""
        _, ext = os.path.splitext(self.file_path)
        self.file_type = ext.lower().replace('.', '')
    
    def _parse_pdf(self):
        """Parse PDF document"""
        text_content = ""
        tables = []
        
        with pdfplumber.open(self.file_path) as pdf:
            for page in pdf.pages:
                # Extract text
                text_content += page.extract_text() + "\n"
                
                # Extract tables
                for table in page.extract_tables():
                    if table:
                        tables.append(table)
        
        # Convert to OCR if text content is minimal
        if len(text_content.strip()) < 100:
            return self._pdf_to_image_ocr()
        
        return {
            'text': text_content,
            'tables': tables
        }
    
    def _pdf_to_image_ocr(self):
        """Convert PDF to images and perform OCR"""
        text_content = ""
        
        # Convert PDF to images
        images = convert_from_path(self.file_path)
        
        # Perform OCR on each image
        for img in images:
            text_content += pytesseract.image_to_string(img) + "\n"
        
        return {
            'text': text_content,
            'tables': []  # OCR doesn't preserve table structure
        }
    
    def _parse_image(self):
        """Parse image using OCR"""
        img = Image.open(self.file_path)
        text_content = pytesseract.image_to_string(img)
        
        # Try to detect tables using OpenCV
        tables = self._detect_tables_in_image()
        
        return {
            'text': text_content,
            'tables': tables
        }
    
    def _detect_tables_in_image(self):
        """Detect tables in image using OpenCV"""
        # This is a simplified implementation
        # A more robust implementation would use line detection and cell extraction
        
        # Read image
        img = cv2.imread(self.file_path)
        if img is None:
            return []
        
        # Convert to grayscale
        gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
        
        # Apply threshold
        _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY_INV)
        
        # Find contours
        contours, _ = cv2.findContours(thresh, cv2.RETR_TREE, cv2.CHAIN_APPROX_SIMPLE)
        
        # Filter contours by size and shape
        table_contours = []
        for contour in contours:
            x, y, w, h = cv2.boundingRect(contour)
            if w > 100 and h > 100 and w/h < 10 and h/w < 10:
                table_contours.append((x, y, w, h))
        
        # For now, just return empty tables
        # In a real implementation, we would extract cells from the detected tables
        return []
    
    def _parse_excel(self):
        """Parse Excel document"""
        # Read all sheets
        excel_data = pd.read_excel(self.file_path, sheet_name=None)
        
        text_content = ""
        tables = []
        
        # Process each sheet
        for sheet_name, df in excel_data.items():
            text_content += f"Sheet: {sheet_name}\n"
            text_content += df.to_string() + "\n\n"
            
            # Convert DataFrame to list of lists for consistency with other parsers
            tables.append(df.values.tolist())
        
        return {
            'text': text_content,
            'tables': tables
        }
    
    def _parse_word(self):
        """Parse Word document"""
        doc = docx.Document(self.file_path)
        
        text_content = ""
        tables = []
        
        # Extract text
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        # Extract tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'text': text_content,
            'tables': tables
        }

def parse_document(file_path):
    """Unified interface for document parsing"""
    parser = DocumentParser()
    return parser.parse(file_path)
